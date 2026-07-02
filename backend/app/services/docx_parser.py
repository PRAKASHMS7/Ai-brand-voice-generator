import docx
import io

class DOCXParser:
    @staticmethod
    def extract_text(file_bytes: bytes) -> str:
        docx_file = io.BytesIO(file_bytes)
        doc = docx.Document(docx_file)
        text_list = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_list.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_list.append(cell.text.strip())
        return "\n".join(text_list)
